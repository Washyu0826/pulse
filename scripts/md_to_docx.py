"""
把 Markdown 轉成 Word .docx（用 python-docx，純程式、不需 Word/pandoc）。

支援：# ## ### 標題、| 表格 |、- / 1. 清單、> 引言、``` 程式碼、**粗體**、`行內碼`。
CJK 用 Microsoft JhengHei。用法：python scripts/md_to_docx.py <a.md> [b.md ...]
"""
import re
import sys

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

_INLINE = re.compile(r"(\*\*.+?\*\*|`[^`]+`)")


def _add_runs(paragraph, text: str) -> None:
    """處理 **粗體** 與 `行內碼`。"""
    pos = 0
    for m in _INLINE.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        tok = m.group()
        if tok.startswith("**"):
            paragraph.add_run(tok[2:-2]).bold = True
        else:  # `code`
            r = paragraph.add_run(tok[1:-1])
            r.font.name = "Consolas"
            r.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def _split_row(line: str) -> list[str]:
    return [c.strip() for c in line.strip().strip("|").split("|")]


def convert(md_path: str, docx_path: str) -> None:
    lines = open(md_path, encoding="utf-8").read().split("\n")
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft JhengHei"
    normal.font.size = Pt(11)
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")

    i, n = 0, len(lines)
    in_code, code_buf = False, []
    while i < n:
        line = lines[i]

        if line.strip().startswith("```"):
            if in_code:
                p = doc.add_paragraph()
                r = p.add_run("\n".join(code_buf))
                r.font.name = "Consolas"
                r.font.size = Pt(9)
            in_code = not in_code
            code_buf = []
            i += 1
            continue
        if in_code:
            code_buf.append(line)
            i += 1
            continue

        # 表格：本行 | ... | 且下一行是 |---|
        if line.strip().startswith("|") and i + 1 < n and re.match(r"^\s*\|[\s:|-]+\|\s*$", lines[i + 1]):
            header = _split_row(line)
            i += 2
            body = []
            while i < n and lines[i].strip().startswith("|"):
                body.append(_split_row(lines[i]))
                i += 1
            table = doc.add_table(rows=1, cols=len(header))
            table.style = "Light Grid Accent 1"
            for j, h in enumerate(header):
                _add_runs(table.rows[0].cells[j].paragraphs[0], h)
            for row in body:
                cells = table.add_row().cells
                for j in range(len(header)):
                    _add_runs(cells[j].paragraphs[0], row[j] if j < len(row) else "")
            doc.add_paragraph()
            continue

        m = re.match(r"^(#{1,6})\s+(.*)$", line)
        if m:
            doc.add_heading(m.group(2).strip(), level=min(len(m.group(1)), 4))
            i += 1
            continue

        if re.match(r"^---+\s*$", line):
            i += 1
            continue

        if line.startswith(">"):
            _add_runs(doc.add_paragraph(style="Intense Quote"), line.lstrip("> ").rstrip())
            i += 1
            continue

        m = re.match(r"^\s*[-*]\s+(.*)$", line)
        if m:
            _add_runs(doc.add_paragraph(style="List Bullet"), m.group(1))
            i += 1
            continue

        m = re.match(r"^\s*\d+\.\s+(.*)$", line)
        if m:
            _add_runs(doc.add_paragraph(style="List Number"), m.group(1))
            i += 1
            continue

        if line.strip() == "":
            i += 1
            continue

        _add_runs(doc.add_paragraph(), line)
        i += 1

    doc.save(docx_path)
    print("saved", docx_path)


if __name__ == "__main__":
    for md in sys.argv[1:]:
        convert(md, md.rsplit(".", 1)[0] + ".docx")
